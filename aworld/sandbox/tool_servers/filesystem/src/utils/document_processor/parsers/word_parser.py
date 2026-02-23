"""Word 解析器"""

import base64
import logging
import zipfile
from pathlib import Path
from typing import Any, Dict, Optional

try:
    from docx import Document
except ImportError:
    Document = None

from ..base_parser import BaseParser

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    def get_supported_types(self) -> list[str]:
        return ["docx", "doc"]

    def can_handle(self, file_type: str) -> bool:
        return file_type.lower() in ["docx", "doc"]

    async def parse(
        self,
        file_path: Path,
        task_id: str = "",
        source_file_id: Optional[str] = None,
        source_file_name: Optional[str] = None,
        afts_service: Any = None,
        output_path: Optional[Path] = None,
        **kwargs,
    ) -> Dict[str, Any]:
        if Document is None:
            raise RuntimeError("未安装 python-docx。请安装: pip install python-docx")
        use_base64 = kwargs.get("use_base64", False)
        out_path = Path(output_path) if output_path else None
        if out_path:
            markdown_content = await self._extract_docx_content(file_path, out_path, use_base64)
            out_path.parent.mkdir(parents=True, exist_ok=True)
            out_path.write_text(markdown_content, encoding="utf-8")
            return {"file_path": out_path}
        source_name = source_file_name or file_path.stem
        parsed_file_path = await self._save_markdown_to_file(
            await self._extract_docx_content(file_path, None, use_base64),
            task_id,
            source_name,
            output_path=output_path,
        )
        return {"file_path": parsed_file_path}

    async def _extract_images_from_docx(
        self, file_path: Path, output_dir: Optional[Path] = None, use_base64: bool = False
    ) -> dict:
        images_map = {}
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                media_files = [f for f in z.namelist() if f.startswith("word/media/")]
                for idx, mf in enumerate(media_files):
                    try:
                        ext = Path(mf).suffix.lower()
                        if ext not in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"]:
                            continue
                        data = z.read(mf)
                        mime = {
                            ".png": "image/png",
                            ".jpg": "image/jpeg",
                            ".jpeg": "image/jpeg",
                            ".gif": "image/gif",
                            ".bmp": "image/bmp",
                            ".tiff": "image/tiff",
                            ".webp": "image/webp",
                        }.get(ext, "image/png")
                        if use_base64:
                            images_map[Path(mf).name] = {
                                "base64": base64.b64encode(data).decode("utf-8"),
                                "mime_type": mime,
                            }
                        elif output_dir:
                            output_dir.mkdir(parents=True, exist_ok=True)
                            img_path = output_dir / f"{file_path.stem}_img_{idx}{ext}"
                            img_path.write_bytes(data)
                            images_map[Path(mf).name] = {"path": str(img_path), "mime_type": mime}
                    except Exception as e:
                        logger.warning("word_parser extract image %s: %s", mf, e)
        except Exception as e:
            logger.warning("word_parser _extract_images_from_docx: %s", e)
        return images_map

    async def _extract_docx_content(
        self, file_path: Path, output_path: Optional[Path] = None, use_base64: bool = False
    ) -> str:
        doc = Document(str(file_path))
        images_dir = None
        if not use_base64 and output_path:
            images_dir = output_path.parent / f"{output_path.stem}_images"
        elif not use_base64:
            images_dir = file_path.parent / f"{file_path.stem}_images"
        images_map = await self._extract_images_from_docx(file_path, images_dir, use_base64)
        image_list = list(images_map.values())
        image_idx = 0
        parts = []
        base_dir = output_path.parent if output_path else file_path.parent
        for para in doc.paragraphs:
            has_image = False
            for run in para.runs:
                if run._element.xpath(".//a:blip"):
                    if image_idx < len(image_list):
                        info = image_list[image_idx]
                        if use_base64:
                            parts.append(
                                f"![图片 {image_idx + 1}](data:{info.get('mime_type', 'image/png')};base64,{info.get('base64', '')})\n\n"
                            )
                        else:
                            p = info.get("path", "")
                            if p and base_dir:
                                rel = Path(p).relative_to(base_dir)
                                parts.append(f"![图片 {image_idx + 1}]({rel})\n\n")
                            elif p:
                                parts.append(f"![图片 {image_idx + 1}]({p})\n\n")
                        image_idx += 1
                        has_image = True
            if para.text.strip():
                style = para.style.name if para.style else "Normal"
                text = para.text.strip()
                if "Heading 1" in style:
                    parts.append(f"# {text}\n\n")
                elif "Heading 2" in style:
                    parts.append(f"## {text}\n\n")
                elif "Heading 3" in style:
                    parts.append(f"### {text}\n\n")
                elif "Heading" in style:
                    parts.append(f"#### {text}\n\n")
                else:
                    parts.append(f"{text}\n\n")
            elif not has_image:
                pass
        if image_idx < len(image_list):
            parts.append("\n## 图片\n\n")
            for i, info in enumerate(image_list[image_idx:], start=image_idx):
                if use_base64:
                    parts.append(
                        f"![图片 {i + 1}](data:{info.get('mime_type', 'image/png')};base64,{info.get('base64', '')})\n\n"
                    )
                else:
                    p = info.get("path", "")
                    if p and base_dir:
                        parts.append(f"![图片 {i + 1}]({Path(p).relative_to(base_dir)})\n\n")
                    elif p:
                        parts.append(f"![图片 {i + 1}]({p})\n\n")
        if doc.tables:
            parts.append("\n## 表格\n\n")
            for ti, table in enumerate(doc.tables):
                parts.append(f"### 表格 {ti + 1}\n\n")
                if table.rows:
                    headers = [c.text.strip() for c in table.rows[0].cells]
                    parts.append("| " + " | ".join(headers) + " |\n")
                    parts.append("|" + "---|" * len(headers) + "\n")
                    for row in table.rows[1:]:
                        parts.append("| " + " | ".join(c.text.strip() for c in row.cells) + " |\n")
                    parts.append("\n")
        return "".join(parts)

